"""
DOCX report export — professional Word document with sections and embedded charts.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any, Dict

from backend.reports.charts import generate_all_charts

logger = logging.getLogger("ai_forge.reports.docx")

REPORTS_DIR = Path("data/forensics/exports")


def export_docx(bundle: Dict[str, Any], output_path: Path) -> str:
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx required: pip install python-docx") from exc

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    meta = bundle.get("meta", {})
    exec_sum = bundle.get("executive_summary", {})
    tech = bundle.get("technical_summary", {})
    evidence = bundle.get("evidence_summary", {})
    evidence_id = meta.get("evidence_id", "unknown")

    charts = generate_all_charts(bundle, evidence_id)
    doc = Document()

    # Title
    title = doc.add_heading("AI-FORGE Digital Forensic Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Report ID: {meta.get('report_id', '—')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Evidence: {evidence_id} | Generated: {meta.get('generated_at', '')[:19]}")
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows_data = [
        ("Verdict", str(exec_sum.get("verdict", "—"))),
        ("Risk Score", f"{exec_sum.get('risk_score', 0):.1f} / 100"),
        ("Risk Level", str(exec_sum.get("risk_level", "—"))),
        ("Confidence", f"{exec_sum.get('confidence', 0):.1f}%"),
    ]
    for i, (k, v) in enumerate(rows_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    if exec_sum.get("narrative"):
        doc.add_paragraph(exec_sum["narrative"])
    doc.add_paragraph(f"Recommendation: {exec_sum.get('recommendation', '—')}")

    doc.add_heading("Key Findings", level=2)
    for f in exec_sum.get("key_findings", []):
        doc.add_paragraph(str(f), style="List Bullet")

    if charts.get("risk_gauge"):
        doc.add_paragraph()
        doc.add_picture(charts["risk_gauge"], width=Inches(3))

    # Technical Summary
    doc.add_page_break()
    doc.add_heading("Technical Summary", level=1)
    doc.add_paragraph(f"Scan mode: {tech.get('scan_mode', 'deep')}")
    if charts.get("module_scores"):
        doc.add_picture(charts["module_scores"], width=Inches(6))

    tampering = tech.get("tampering") or {}
    if tampering:
        doc.add_heading("Tampering Analysis", level=2)
        doc.add_paragraph(
            f"Verdict: {tampering.get('verdict')} | Severity: {tampering.get('severity')} | "
            f"Score: {tampering.get('score')}"
        )

    # Evidence Summary
    if evidence.get("registered"):
        doc.add_page_break()
        doc.add_heading("Evidence Summary", level=1)
        ev_table = doc.add_table(rows=6, cols=2)
        ev_table.style = "Table Grid"
        ev_data = [
            ("Filename", evidence.get("original_filename", "—")),
            ("SHA-256", evidence.get("sha256", "—")),
            ("SHA-512", (evidence.get("sha512") or "—")[:64]),
            ("Intake", str(evidence.get("intake_timestamp", ""))[:19]),
            ("Custody Events", str(evidence.get("custody_events", 0))),
            ("Chain Verified", "YES" if evidence.get("chain_verified") else "NO"),
        ]
        for i, (k, v) in enumerate(ev_data):
            ev_table.rows[i].cells[0].text = k
            ev_table.rows[i].cells[1].text = str(v)

    # Timeline
    timeline = bundle.get("timeline", [])
    if timeline:
        doc.add_page_break()
        doc.add_heading("Investigation Timeline", level=1)
        if charts.get("timeline"):
            doc.add_picture(charts["timeline"], width=Inches(6))
        for event in timeline[:25]:
            ts = (event.get("timestamp") or "")[:19]
            doc.add_paragraph(f"[{ts}] {event.get('type')}: {event.get('description')}", style="List Bullet")

    # Court
    court = bundle.get("court_certification")
    if court:
        doc.add_page_break()
        doc.add_heading("Court Report Certification", level=1)
        doc.add_paragraph(court.get("certification", ""))
        doc.add_paragraph(court.get("chain_of_custody_attestation", ""))

    # Heatmaps
    artifacts = bundle.get("artifacts") or {}
    vis_keys = [k for k in artifacts if k in ("ela", "edges", "wavelet", "copy_move", "heatmap") or "heatmap" in k]
    if vis_keys:
        doc.add_page_break()
        doc.add_heading("Forensic Visualizations", level=1)
        for key in vis_keys[:4]:
            path = artifacts[key]
            if Path(path).exists():
                doc.add_heading(key.replace("_", " ").title(), level=2)
                try:
                    doc.add_picture(path, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(f"[Image: {path}]")

    doc.add_paragraph()
    p = doc.add_paragraph(
        "This report is AI-assisted and should be verified by a qualified forensic expert."
    )
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    doc.save(str(output_path))
    return str(output_path)
